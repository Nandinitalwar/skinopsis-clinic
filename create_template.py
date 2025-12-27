#!/usr/bin/env python3
"""
Script to create a basic Word template based on the user's format
"""
from docx import Document
from docx.shared import Inches

def create_template():
    doc = Document()
    
    # Header section
    header = doc.add_paragraph()
    header.add_run("Patient Name: {{patient_name}}\n")
    header.add_run("Age: {{age_years}} years\n")
    header.add_run("Sex: {{sex}}\n\n")
    
    # Diagnosis section
    diagnosis = doc.add_paragraph()
    diagnosis.add_run("Diagnosis:\n{{diagnosis}}\n\n")
    
    # Duration section
    duration = doc.add_paragraph()
    duration.add_run("Duration of Symptoms:\n{{symptom_duration}}\n\n")
    
    # Presenting symptoms
    symptoms = doc.add_paragraph()
    symptoms.add_run("Presenting Symptoms:\n{{presenting_symptoms_block}}\n\n")
    
    # Allergies
    allergies = doc.add_paragraph()
    allergies.add_run("Allergies:\n{{allergies}}\n\n")
    
    # Current medications
    current_meds = doc.add_paragraph()
    current_meds.add_run("Current Medications:\n{{current_medications}}\n\n")
    
    # Past medical history
    history = doc.add_paragraph()
    history.add_run("Past Medical History:\n{{past_medical_history}}\n\n")
    
    # Treatment plan
    treatment = doc.add_paragraph()
    treatment.add_run("Treatment Plan:\n{{treatment_plan_block}}\n\n")
    
    # Follow-up
    followup = doc.add_paragraph()
    followup.add_run("{{followup_text}}\n\n")
    
    # Footer with doctor info
    footer = doc.add_paragraph()
    footer.add_run("Sincerely,\n\n")
    footer.add_run("{{date}}\n")
    footer.add_run("Dr. Amrita Kaur Talwar\n")
    footer.add_run("Consultant Dermatologist\n")
    footer.add_run("M.B.B.S, M.D., D.D (U.K)\n")
    footer.add_run("Medical License Number : MMC : 2012/10/2891\n")
    footer.add_run("402 - 407 Doctor Centre, 135,August Kranti Marg, Mumbai - 400036.\n")
    footer.add_run("402clinic@gmail.com | +912223632524\n")
    
    # Save the document
    doc.save("templates/prescription_template.docx")
    print("Template created successfully at templates/prescription_template.docx")

if __name__ == "__main__":
    create_template()