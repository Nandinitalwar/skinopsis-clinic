#!/usr/bin/env python3
"""
Script to update the Word template with better spacing and signature
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_improved_template():
    doc = Document()
    
    # Set smaller margins for more space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header section - more compact
    header = doc.add_paragraph()
    header_run = header.add_run("Patient Name: {{patient_name}}    Age: {{age_years}} years    Sex: {{sex}}")
    header_run.font.size = Pt(11)
    header_run.bold = True
    header.space_after = Pt(6)
    
    # Diagnosis section - compact
    diagnosis = doc.add_paragraph()
    diagnosis_run = diagnosis.add_run("Diagnosis: {{diagnosis}}")
    diagnosis_run.font.size = Pt(11)
    diagnosis_run.bold = True
    diagnosis.space_after = Pt(4)
    
    # Duration section - compact
    duration = doc.add_paragraph()
    duration_run = duration.add_run("Duration of Symptoms: {{symptom_duration}}")
    duration_run.font.size = Pt(10)
    duration.space_after = Pt(4)
    
    # Presenting symptoms - compact
    symptoms = doc.add_paragraph()
    symptoms_title = symptoms.add_run("Presenting Symptoms:")
    symptoms_title.font.size = Pt(10)
    symptoms_title.bold = True
    symptoms.add_run("\n{{presenting_symptoms_block}}")
    symptoms.space_after = Pt(6)
    
    # Allergies, medications, history in a more compact format
    allergies = doc.add_paragraph()
    allergies_run = allergies.add_run("Allergies: {{allergies}}")
    allergies_run.font.size = Pt(10)
    allergies.space_after = Pt(4)
    
    current_meds = doc.add_paragraph()
    current_meds_run = current_meds.add_run("Current Medications: {{current_medications}}")
    current_meds_run.font.size = Pt(10)
    current_meds.space_after = Pt(4)
    
    history = doc.add_paragraph()
    history_run = history.add_run("Past Medical History: {{past_medical_history}}")
    history_run.font.size = Pt(10)
    history.space_after = Pt(6)
    
    # Treatment plan - prominent but compact
    treatment = doc.add_paragraph()
    treatment_title = treatment.add_run("Treatment Plan:")
    treatment_title.font.size = Pt(11)
    treatment_title.bold = True
    treatment.add_run("\n{{treatment_plan_block}}")
    treatment.space_after = Pt(8)
    
    # Follow-up
    followup = doc.add_paragraph()
    followup_run = followup.add_run("{{followup_text}}")
    followup_run.font.size = Pt(10)
    followup.space_after = Pt(12)
    
    # Signature section - compact footer
    signature_section = doc.add_paragraph()
    signature_section.add_run("Sincerely,").font.size = Pt(10)
    
    # Add space for signature image placeholder
    sig_para = doc.add_paragraph()
    sig_para.add_run("Dr. Amrita Kaur Talwar").bold = True
    sig_para.space_before = Pt(24)  # Space for signature
    
    # Doctor details - compact
    details = doc.add_paragraph()
    details_run = details.add_run(
        "Consultant Dermatologist\n"
        "M.B.B.S, M.D., D.D (U.K)\n"
        "Medical License Number: MMC: 2012/10/2891\n"
        "402 - 407 Doctor Centre, 135, August Kranti Marg, Mumbai - 400036\n"
        "402clinic@gmail.com | +912223632524"
    )
    details_run.font.size = Pt(9)
    
    # Date at the end
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run("{{date}}")
    date_run.font.size = Pt(10)
    
    # Save the document
    doc.save("templates/prescription_template.docx")
    print("Updated template created successfully at templates/prescription_template.docx")

if __name__ == "__main__":
    create_improved_template()