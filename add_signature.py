#!/usr/bin/env python3
"""
Script to add signature to the Word template
Place your signature image as 'signature.png' in the templates folder
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_signature_to_template():
    # Load existing template
    doc = Document("templates/prescription_template.docx")
    
    # Find the signature paragraph (look for "Sincerely")
    for i, para in enumerate(doc.paragraphs):
        if "Sincerely" in para.text:
            # Add a new paragraph for the signature image
            sig_para = doc.paragraphs[i+1]
            
            # Check if signature image exists
            signature_path = "templates/signature.png"
            if os.path.exists(signature_path):
                # Clear the paragraph and add signature image
                sig_para.clear()
                run = sig_para.add_run()
                run.add_picture(signature_path, width=Inches(2.0))  # Adjust width as needed
                
                # Add doctor name below signature
                doc.paragraphs[i+1].add_run("\nDr. Amrita Kaur Talwar").bold = True
            else:
                print(f"Signature image not found at {signature_path}")
                print("Please place your signature image as 'signature.png' in the templates folder")
                return False
            break
    
    # Save the updated template
    doc.save("templates/prescription_template.docx")
    print("Signature added to template successfully!")
    return True

if __name__ == "__main__":
    add_signature_to_template()